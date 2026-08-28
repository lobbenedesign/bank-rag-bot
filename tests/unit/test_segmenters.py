"""Round-trip tests for the per-format segmenters — real bytes in, real
DocumentSegments with correct locators out. Not mocked: every format here,
including PDF, is verified against files actually built with a real writer
library (python-docx, openpyxl, fpdf2), so a broken segmenter fails here,
not silently in production.

fpdf2 is a dev/test-only dependency (see pyproject.toml's [dev] extra) —
never imported by application code, only used here to generate a real
text-bearing PDF with controlled paragraph gaps for segment_pdf to parse.
"""
from __future__ import annotations

import io
import json

from bank_rag.ingestion.segmentation.csv_segmenter import segment_csv
from bank_rag.ingestion.segmentation.docx_segmenter import segment_docx
from bank_rag.ingestion.segmentation.json_segmenter import segment_json
from bank_rag.ingestion.segmentation.markdown_segmenter import segment_markdown, segment_plain_text
from bank_rag.ingestion.segmentation.pdf_segmenter import segment_pdf
from bank_rag.ingestion.segmentation.xlsx_segmenter import segment_xlsx
from bank_rag.ingestion.segmentation.xml_segmenter import segment_xml


def test_markdown_segments_by_heading_with_breadcrumb_locator():
    md = (
        "# Conti correnti\n"
        "Testo introduttivo.\n"
        "## Conto Base\n"
        "Il Conto Base non ha canone mensile.\n"
        "## Conto Business\n"
        "Il Conto Business ha canone di 5 euro al mese.\n"
    )
    segments = segment_markdown(md.encode("utf-8"))

    assert [s.locator.value for s in segments] == [
        "Conti correnti",
        "Conti correnti > Conto Base",
        "Conti correnti > Conto Business",
    ]
    assert all(s.locator.kind == "section" for s in segments)
    assert "canone mensile" in segments[1].text


def test_plain_text_falls_back_to_line_ranges():
    text = "\n".join(f"riga {i}" for i in range(1, 101))
    segments = segment_plain_text(text.encode("utf-8"))

    assert segments[0].locator.kind == "line_range"
    assert segments[0].locator.value == "1-40"
    assert segments[1].locator.value == "41-80"
    assert segments[2].locator.value == "81-100"


def test_markdown_without_headings_falls_back_to_plain_text():
    text = "\n".join(f"riga {i}" for i in range(1, 50))
    segments = segment_markdown(text.encode("utf-8"))
    assert segments[0].locator.kind == "line_range"


def test_csv_segments_in_row_batches_with_header_context():
    rows = ["nome,tasso"] + [f"prodotto{i},{i}.5" for i in range(1, 61)]
    csv_bytes = "\n".join(rows).encode("utf-8")

    segments = segment_csv(csv_bytes)

    assert len(segments) == 2  # 60 data rows / 50 per segment
    assert segments[0].locator.kind == "row_range"
    assert segments[0].locator.value == "2-51"
    assert segments[1].locator.value == "52-61"
    assert "nome, tasso" in segments[0].text  # header repeated as context
    assert "prodotto1, 1.5" in segments[0].text
    assert "prodotto60, 60.5" in segments[1].text


def test_json_object_root_segments_per_top_level_key():
    data = {"conto_base": {"canone": 0}, "conto_business": {"canone": 5}}
    segments = segment_json(json.dumps(data).encode("utf-8"))

    assert {s.locator.value for s in segments} == {"$.conto_base", "$.conto_business"}
    assert all(s.locator.kind == "json_path" for s in segments)


def test_json_array_root_segments_per_item():
    data = [{"id": 1}, {"id": 2}, {"id": 3}]
    segments = segment_json(json.dumps(data).encode("utf-8"))

    assert [s.locator.value for s in segments] == ["$[0]", "$[1]", "$[2]"]


def test_xml_segments_per_root_child_with_indexed_xpath():
    xml = (
        "<products>"
        "<product><name>Conto Base</name></product>"
        "<product><name>Conto Business</name></product>"
        "<promo><name>Estate 2026</name></promo>"
        "</products>"
    )
    segments = segment_xml(xml.encode("utf-8"))

    assert [s.locator.value for s in segments] == [
        "/products/product[1]",
        "/products/product[2]",
        "/products/promo[1]",
    ]
    assert all(s.locator.kind == "xpath" for s in segments)
    assert "Conto Base" in segments[0].text


def test_docx_segments_by_heading_section():
    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_heading("Conti correnti", level=1)
    document.add_paragraph("Testo introduttivo.")
    document.add_heading("Conto Base", level=2)
    document.add_paragraph("Il Conto Base non ha canone mensile.")
    document.add_heading("Conto Business", level=2)
    document.add_paragraph("Il Conto Business ha canone di 5 euro.")

    buffer = io.BytesIO()
    document.save(buffer)

    segments = segment_docx(buffer.getvalue())

    assert [s.locator.value for s in segments] == [
        "Conti correnti",
        "Conti correnti > Conto Base",
        "Conti correnti > Conto Business",
    ]
    assert "canone mensile" in segments[1].text


def test_xlsx_segments_per_sheet_row_batches():
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Tassi"
    sheet.append(["prodotto", "tasso"])
    for i in range(1, 61):
        sheet.append([f"prodotto{i}", i * 0.1])

    buffer = io.BytesIO()
    workbook.save(buffer)

    segments = segment_xlsx(buffer.getvalue())

    assert len(segments) == 2
    assert segments[0].locator.kind == "row_range"
    assert segments[0].locator.value == "Tassi!2-51"
    assert segments[1].locator.value == "Tassi!52-61"


def test_pdf_segments_by_page_and_paragraph():
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, "Primo paragrafo della pagina uno, con del testo di condizioni bancarie.")
    pdf.ln(20)  # large gap: must be read as a paragraph break, not a line break
    pdf.multi_cell(0, 8, "Secondo paragrafo della pagina uno, con altre condizioni.")
    pdf.add_page()
    pdf.multi_cell(0, 8, "Primo paragrafo della pagina due.")

    content = bytes(pdf.output())

    segments = segment_pdf(content)

    assert all(s.locator.kind == "page_paragraph" for s in segments)
    pages = {s.locator.value.split(":")[0] for s in segments}
    assert pages == {"1", "2"}

    page_one_paragraphs = [s.text for s in segments if s.locator.value.startswith("1:")]
    assert len(page_one_paragraphs) == 2  # the large ln(20) gap split into two paragraphs
    assert "Primo paragrafo della pagina uno" in page_one_paragraphs[0]
    assert "Secondo paragrafo della pagina uno" in page_one_paragraphs[1]

    page_two_paragraphs = [s.text for s in segments if s.locator.value.startswith("2:")]
    assert any("Primo paragrafo della pagina due" in text for text in page_two_paragraphs)


def test_pdf_with_no_extractable_text_yields_no_segments():
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()  # blank page, no text drawn
    content = bytes(pdf.output())

    assert segment_pdf(content) == []


def test_pdf_table_becomes_a_real_markdown_table_not_flattened_prose():
    """A rate table read left-to-right as prose loses the row/column
    alignment between a duration and its rate — the exact failure mode a
    banking rate sheet cannot tolerate. This exercises the real fix: a
    detected table (fpdf2's own bordered `pdf.table()`, which draws real
    ruling lines pdfplumber's find_tables() can detect) becomes its own
    Markdown-table segment, and the table's own text is excluded from the
    surrounding paragraph text (not duplicated, not garbled).
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "Foglio tassi mutuo giovani", new_x="LMARGIN", new_y="NEXT")
    with pdf.table() as table:
        header = table.row()
        for h in ("Durata", "Tasso Fisso", "Tasso Variabile"):
            header.cell(h)
        for values in (("20 anni", "3.25%", "2.90%"), ("30 anni", "3.50%", "3.10%")):
            row = table.row()
            for v in values:
                row.cell(v)
    content = bytes(pdf.output())

    segments = segment_pdf(content)

    table_segments = [s for s in segments if s.locator.kind == "page_table"]
    assert len(table_segments) == 1
    assert table_segments[0].locator.value == "1:1"
    table_text = table_segments[0].text
    # A real Markdown table: header row, separator row, one row per data row —
    # not "Durata Tasso Fisso Tasso Variabile 20 anni 3.25% 2.90% ..." mashed together.
    assert table_text.splitlines()[0] == "| Durata | Tasso Fisso | Tasso Variabile |"
    assert "| --- | --- | --- |" in table_text
    assert "| 20 anni | 3.25% | 2.90% |" in table_text
    assert "| 30 anni | 3.50% | 3.10% |" in table_text

    # The table's own cell text must not also leak into a paragraph segment
    # (would mean the exclusion-by-bbox logic isn't working).
    paragraph_segments = [s for s in segments if s.locator.kind == "page_paragraph"]
    assert not any("3.25%" in s.text for s in paragraph_segments)
    assert any("Foglio tassi mutuo giovani" in s.text for s in paragraph_segments)
